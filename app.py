import os, time, re, base64, io
import streamlit as st
from huggingface_hub import InferenceClient
from httpx import ConnectTimeout, ReadTimeout, HTTPError
from docx import Document

# Hugging Face 設定
HF_TOKEN = os.getenv("HUGGINGFACEHUB_API_TOKEN", "")
MODEL_ID = "Qwen/Qwen2.5-7B-Instruct"

# CSS
def load_css(path: str):
    with open(path, "r", encoding="utf-8") as f:
        css = f.read()
    st.markdown(f"<style>{css}</style>", unsafe_allow_html=True)

def load_base64_image(path):
    with open(path, "rb") as f:
        data = f.read()
    return base64.b64encode(data).decode()

logo_light = load_base64_image("img/logo_black.PNG")
logo_dark = load_base64_image("img/logo_white.PNG")

# セッション初期化
def init_session_state():
    st.session_state.setdefault("current_ad", "")
    st.session_state.setdefault("edited_ad", "")
    st.session_state.setdefault("current_char_count", 0)

# 前処理
def preprocess(text: str) -> str:
    pattern = re.compile(r'【.*?】|[ＲR][ー-]\d+|■|＊')
    return pattern.sub('', text or "")

# 後処理
def postprocess(text: str) -> str:
    text = (text or "").strip()
    text = re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL)
    return text.replace("\n", "").replace("\r", "").strip()

# 文字数
def count(text: str) -> int:
    return len((text or "").replace("\n", "").replace("\r", ""))

def realtime_count():
    text = st.session_state.get("edited_ad", "")
    st.session_state["current_char_count"] = count(text)

# Word保存
def save_docx(text):
    doc = Document()
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# HF 呼び出し
def _hf_text(choice) -> str:
    msg = getattr(choice, "message", None)
    if isinstance(msg, dict):
        return msg.get("content", "")
    return getattr(msg, "content", "")

def _hf_chat(client, messages, max_tokens, temperature):
    for attempt in range(3):
        try:
            resp = client.chat.completions.create(
                messages=messages,
                max_tokens=max_tokens,
                temperature=temperature,
            )
            return _hf_text(resp.choices[0]).strip()
        except (ConnectTimeout, ReadTimeout):
            time.sleep(2 ** attempt)
        except HTTPError as e:
            if getattr(e.response, "status_code", 500) >= 500:
                time.sleep(2 ** attempt)
            else:
                raise
    raise RuntimeError("HF API から応答を取得できませんでした")

# 文字数厳格化
def adjust_length(client, ad, target_length, tone, max_tokens, temperature):
    ad = postprocess(ad)
    tone = build_tone(tone)
    system_prompt = "あなたは日本語文章の文字数を正確に調整する編集者です。"

    for _ in range(2):
        current_length = len(ad)
        sub = target_length - current_length

        if abs(sub) <= 5 and ad.endswith("。"):
            return ad

        if sub > 0:
            sub_length = f"現在 {current_length}文字なので、{sub} 文字分だけ内容を自然に補ってください。"
        else:
            sub_length = f"現在 {current_length}文字なので、{abs(sub)} 文字分だけ内容を自然に削ってください。"

        prompt = (
            f"次の文章の意味と構成をできるだけ変えずに、文字数だけを調整してください。\n"
            f"{tone}"
            f"{sub_length}\n"
            f"文字数が {target_length} ±5 の範囲に入っていない場合は、必ず文章を修正して再生成\n"
            f"生成文のみを出力・捕捉などは出力しない\n"
            f"【元の文章】\n{ad}\n\n"
            f"【整形後（{target_length}文字）】"
        )

        new_ad = _hf_chat(
            client,
            [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": prompt},
            ],
            max_tokens=max_tokens,
            temperature=temperature,
        )

        if not new_ad:
            break

        ad = postprocess(new_ad)

    return ad

# キーワード指定
def split_keywords(keywords: str):
    return [w for w in re.split(r"[ 　]+", keywords.strip()) if w]

def build_keyword(keywords: str):
    words = split_keywords(keywords)
    if not words:
        return ""
    return (
        f"・キーワード指定: { '、'.join(words) }\n"
        "・各キーワードは文章中に絶対に1回だけ\n"
        "・キーワードが文の繋がりを邪魔しないよう、自然に組み込んでください。\n"
    )

# 表現
def build_tone(tone: str) -> str:
    if tone == "やわらかい":
        return (
            "・話し言葉に近い文体で書くこと\n"
            "・一文を短く、簡単な言葉で書くこと\n"
            "・難しい言い回しや抽象語は使わないこと\n"
        )
    else:
        return ""

# 広告文生成
def generate_advertisement(text, target_length, keywords, tone, temperature=0.2):
    if not HF_TOKEN:
        raise RuntimeError("HUGGINGFACEHUB_API_TOKEN が設定されていません。")

    client = InferenceClient(model=MODEL_ID, token=HF_TOKEN, timeout=60)
    cleaned = preprocess(text)
    max_tokens=int((target_length) * 1.2)
    tone = build_tone(tone)
    keyword = build_keyword(keywords)

    ad = _hf_chat(
        client,
        [{"role": "user", "content": (
            f"次の原稿をもとに、新聞のラテ欄風の広告文を書いてください。\n"
            f"文字数は日本語でちょうど {target_length} 文字\n"
            f"次にある文体ルールを厳密に守ってください。\n"
            f"・文の主語・場所・行動が明確になるように書くこと\n"
            f"・文末には視聴者の興味を引くフックを必ず入れる\n"
            f"・感情やインパクトのある言葉を使う\n"
            f"・季節感やテーマがあれば冒頭に入れる\n"
            f"・応募やプレゼントキャンペーンについては書かない\n"
            f"【条件】日本語のみ・改行なし\n"
            f"{tone}\n"
            f"{keyword}\n\n"
            f"【原稿】\n{cleaned}\n\n【広告文】"
        )}],
        max_tokens=max_tokens,
        temperature=temperature,
    )

    ad = postprocess(ad)
    ad = adjust_length(client, ad, target_length, tone, max_tokens, temperature=0.1)

    return ad

# Streamlit UI
def main():
    st.set_page_config(page_title="南海ことば工房", page_icon="img/favicon.JPG")
    load_css("style.css")
    init_session_state()
    
    st.markdown(
        f'<div class="logo-container"><img src="data:image/png;base64,{logo_light}" class="logo-light"><img src="data:image/png;base64,{logo_dark}" class="logo-dark"></div>', unsafe_allow_html=True)

    option = st.sidebar.radio("入力方法を選択", ("テキスト", "ファイル"))
    text = ""

    if option == "テキスト":
        text = st.text_area("広告文にしたい原稿を入力してください", height=260)
        text = text.strip().replace("\r", "")
        text = text.replace("　", " ")
        text = re.sub(r"\s+", " ", text)

    else:
        uploadfile = st.file_uploader("ファイルを選択", type=["txt", "docx"])
        if uploadfile is not None:
            try:
                if uploadfile.name.endswith(".txt"):
                    text = uploadfile.read().decode("utf-8", errors="ignore")
                elif uploadfile.name.endswith(".docx"):
                    doc = Document(uploadfile)
                    text = "\n".join([p.text for p in doc.paragraphs])
            except Exception as e:
                st.error(f"ファイルの読み込みに失敗しました: {e}")
                text = ""

    # 文字数指定
    target_length = st.sidebar.number_input("文字数", min_value=10, max_value=500, value=100, step=1)
    
    # 文章表現選択
    tone = st.sidebar.radio("文章のスタイル", ["かたい", "やわらかい"], horizontal=True)
    
    # キーワード指定
    keywords = st.sidebar.text_input("キーワード指定（スペース区切り）", value="")

    # 保存ファイル名
    filename = st.sidebar.text_input("保存するファイル名", value="newspaper")
    ext = st.sidebar.radio("保存形式", [".txt", ".docx"], horizontal=True)
    download = filename + ext

    # 生成    
    if st.button("広告文を生成"):
        try:
            if not text.strip():
                st.warning("原稿を入力してください。")
            else:
                with st.spinner("広告文を生成中..."):
                    ad = generate_advertisement(text=text, target_length=target_length, keywords=keywords, tone=tone)

                    st.session_state["current_ad"] = ad
                    st.session_state["edited_ad"] = ad
                    st.session_state["current_char_count"] = len(ad)
       
        except Exception as e:
            st.error(f"エラーが発生しました: {e}")

    # ダウンロードボタン
    if st.session_state["current_ad"]:
        st.text_area("生成結果", key="edited_ad", on_change=realtime_count, height=200)
        st.markdown(f"文字数：{st.session_state['current_char_count']} 文字")
                
        final_text = st.session_state["edited_ad"]
        if ext == ".docx":
            file_data = save_docx(final_text)
            mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            file_data = final_text
            mime = "text/plain"
  
        st.download_button(
            label="ダウンロード",
            data=file_data,
            file_name=download,
            mime=mime
            )

# 実行
if __name__ == "__main__":
    main()
