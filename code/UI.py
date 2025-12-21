import os, time, re, base64
import streamlit as st
from huggingface_hub import InferenceClient
from httpx import ConnectTimeout, ReadTimeout, HTTPError
from docx import Document
import io

# ------------------------------------------------
# Hugging Face 設定
# ------------------------------------------------
HF_TOKEN = os.getenv("HUGGINGFACEHUB_API_TOKEN", "")
MODEL_ID = "Qwen/Qwen3-4B-Instruct-2507"

# =========================
# 前処理
# =========================
def remove_strings(text: str) -> str:
    pattern = re.compile(r'【.*?】|[ＲR][ー-]\d+|\n|\t|\s+|■|＊')
    return pattern.sub('', text or "")

def normalize_output(text: str) -> str:
    text = (text or "").strip()
    text = re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL)
    text = re.sub(r"[a-zA-Z]+", "", text)
    return text.replace("\n", "").replace("\r", "").strip()

def count_chars(text: str) -> int:
    return len((text or "").replace("\n", "").replace("\r", ""))

# =========================
# 文分割
# =========================
def split_sentences(text: str) -> list[str]:
    text = normalize_output(text)
    parts = re.split(r"(。)", text)
    sentences = []
    for i in range(0, len(parts) - 1, 2):
        sentences.append(parts[i] + "。")
    return sentences

# =========================
# HF 応答抽出 & 呼び出し
# =========================
def _extract_message_text(choice) -> str:
    msg = getattr(choice, "message", None)
    if isinstance(msg, dict):
        return msg.get("content", "")
    return getattr(msg, "content", "")

def _call_chat(client, messages, max_tokens, temperature):
    for attempt in range(3):
        try:
            resp = client.chat.completions.create(
                messages=messages,
                max_tokens=max_tokens,
                temperature=temperature,
            )
            return _extract_message_text(resp.choices[0]).strip()

        except (ConnectTimeout, ReadTimeout):
            time.sleep(2 ** attempt)
        except HTTPError as e:
            if getattr(e.response, "status_code", 500) >= 500:
                time.sleep(2 ** attempt)
            else:
                raise
    return ""

# =========================
# 後半だけ再調整
# =========================
def adjust_tail_with_llm(
    client,
    head_sentences: list[str],
    tail_sentence: str,
    target_chars: int,
    temperature: float,
) -> str:
    head = "".join(head_sentences)

    prompt = (
        f"次の新聞広告文の【後半】だけを、意味を変えずに自然に書き直してください。\n"
        f"【条件】\n"
        f"・前半は変更しない\n"
        f"・全体がだいたい {target_chars} 文字前後になるよう調整\n"
        f"・改行なし一段落\n"
        f"・固有名詞を使わない\n"
        f"・誇大表現を避ける\n"
        f"・文末は必ず「。」で終える\n\n"
        f"【前半】\n{head}\n\n"
        f"【後半（修正対象）】\n{tail_sentence}\n\n"
        f"【修正後の後半】"
    )

    new_tail = _call_chat(
        client,
        [{"role": "user", "content": prompt}],
        max_tokens=200,
        temperature=temperature,
    )

    return head + normalize_output(new_tail)

# =========================
# 広告文生成（精度優先）
# =========================
def generate_newspaper_ad_api(text: str, target_chars: int, temperature: float = 0.2) -> str:
    if not HF_TOKEN:
        raise RuntimeError("HUGGINGFACEHUB_API_TOKEN が設定されていません。")

    client = InferenceClient(model=MODEL_ID, token=HF_TOKEN, timeout=60.0)

    cleaned = remove_strings(text)
    max_tokens = int(target_chars * 3)

    # ① まず自然さ最優先で生成
    ad = _call_chat(
        client,
        [{"role": "user", "content": (
            f"次の原稿内容をもとに、新聞に掲載できる広告文を作成してください。\n"
            f"・日本語のみ\n"
            f"・固有名詞を使わない\n"
            f"・改行なし一段落\n"
            f"・文末は必ず「。」\n"
            f"・文字数はおよそ {target_chars} 文字前後\n"
            f"・無理な文字数合わせはしない\n\n"
            f"【原稿】\n{cleaned}\n\n【広告文】"
        )}],
        max_tokens=max_tokens,
        temperature=temperature,
    )

    ad = normalize_output(ad)
    sentences = split_sentences(ad)

    # ② 長すぎる場合のみ「最後の1文」を再調整
    if len(sentences) >= 2 and len(ad) > target_chars + 10:
        ad = adjust_tail_with_llm(
            client,
            head_sentences=sentences[:-1],
            tail_sentence=sentences[-1],
            target_chars=target_chars,
            temperature=temperature,
        )

    return ad


# keyword指定のスペース区切りをリスト化する関数
def build_keyword_instruction(keywords: str):
    if not keywords.strip():
        return ""

    # スペース区切り → リスト化
    words = [w.strip() for w in keywords.split() if w.strip()]

    if not words:
        return ""

    joined = "、".join(words)
    return f"以下のキーワードを必ず含めてください：{joined}。"

# wordダウンロードできるように
def create_docx_bytes(text: str):
    doc = Document()
    doc.add_paragraph(text)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# =========================
# Streamlit UI
# =========================
def main():
    
    if "current_ad" not in st.session_state:
        st.session_state["current_ad"] = ""

    # 履歴保持
    if "history" not in st.session_state:
        st.session_state["history"] = []

    st.title("要約")
    
    option = st.sidebar.radio("入力方法を選択", ("テキスト", "ファイル"))
    text = ""

    # 入力方法：テキスト
    if option == "テキスト":
        text = st.text_area("広告文にしたい原稿を入力してください", height=260)

    # 入力方法：ファイル
    else:
        uploadfile = st.file_uploader("ファイルを選択", type=["txt", "docx"])
        if uploadfile is not None:
            try:
                # ---- TXT ----
                if uploadfile.name.endswith(".txt"):
                    text = uploadfile.read().decode("utf-8", errors="ignore")

                # ---- DOCX ----
                elif uploadfile.name.endswith(".docx"):
                    doc = Document(uploadfile)
                    text = "\n".join([p.text for p in doc.paragraphs])

                # st.success("ファイルを読み込みました")

            except Exception as e:
                st.error(f"ファイルの読み込みに失敗しました: {e}")
                text = ""

    # 文字数指定
    target_chars = st.sidebar.number_input(
        "文字数",
        min_value=10,
        max_value=500,
        value=120,
        step=1
    )
    
    # 文章表現選択
    tone = st.sidebar.selectbox(
        "文章の硬さを選択してください",
        ["やわらかい", "ふつう", "かたい"]
        )
    
    # キーワード指定
    keywords = st.sidebar.text_input(
        "キーワード指定（スペース区切り）",
        value=""
        )

    # 保存ファイル名
    filename = st.sidebar.text_input(
        "保存するファイル名",
        value="newspaper"
    )
    ext = st.sidebar.selectbox( "保存形式", [".txt", ".docx"] )
    download = filename + ext
    
    # 履歴    
    st.sidebar.subheader("履歴")
    for i, item in enumerate(st.session_state["history"], start=1):
        with st.sidebar.expander(f"履歴{i}"):
            st.write(item)

    # =========================
    # 要約生成
    # =========================
    if st.button("広告文を生成"):
        try:
            if not text.strip():
                st.warning("原稿を入力してください。")
            else:
                with st.spinner("広告文を生成中..."):
                    ad = generate_newspaper_ad_api(text, target_chars)
                    st.session_state["current_ad"] = ad
                    st.session_state["history"].insert(0, ad)
                    st.session_state["history"] = st.session_state["history"][:5]
                
        except Exception as e:
            st.error(f"エラーが発生しました: {e}")

    # ダウンロードボタン（名前指定）
    if st.session_state["current_ad"]:
        st.text_area(
            "生成された広告文",
            st.session_state["current_ad"],
            height=200
            )
        st.markdown(f"文字数：{len(st.session_state['current_ad'])} 文字")
                
        if ext == ".docx":
            file_data = create_docx_bytes(st.session_state["current_ad"])
            mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            file_data = st.session_state["current_ad"]
            mime = "text/plain"
            
        st.download_button(
            label="📥 ダウンロード",
            data=file_data,
            file_name=download,
            mime=mime
            )

# =========================
# 実行
# =========================
if __name__ == "__main__":
    main()
