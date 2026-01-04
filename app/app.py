import os, time, re, base64, io
import streamlit as st
from huggingface_hub import InferenceClient
from httpx import ConnectTimeout, ReadTimeout, HTTPError
from docx import Document

# Hugging Face 設定
HF_TOKEN = os.getenv("HUGGINGFACEHUB_API_TOKEN", "")
MODEL_ID = "Qwen/Qwen3-4B-Instruct-2507"

# CSS
def load_css(path: str):
    with open(path, "r", encoding="utf-8") as f:
        css = f.read()
    st.markdown(f"<style>{css}</style>", unsafe_allow_html=True)

# 前処理
def remove_strings(text: str) -> str:
    pattern = re.compile(r'【.*?】|[ＲR][ー-]\d+|\n|\t|\s+|■|＊')
    return pattern.sub('', text or "")

def normalize_output(text: str) -> str:
    text = (text or "").strip()
    text = re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL)
    #text = re.sub(r"[a-zA-Z]+", "", text)
    return text.replace("\n", "").replace("\r", "").strip()

def count_chars(text: str) -> int:
    return len((text or "").replace("\n", "").replace("\r", ""))

# HF 呼び出し
def _extract_message_text(choice) -> str:
    msg = getattr(choice, "message", None)
    if isinstance(msg, dict):
        return msg.get("content", "")
    return getattr(msg, "content", "")

def _call_chat(client, messages, max_tokens, temperature):
    for attempt in range(3):
        try:
            resp = client.chat.completions.create(
                messages=messages,
                max_tokens=max_tokens,
                temperature=temperature,
            )
            return _extract_message_text(resp.choices[0]).strip()
        except (ConnectTimeout, ReadTimeout):
            time.sleep(2 ** attempt)
        except HTTPError as e:
            if getattr(e.response, "status_code", 500) >= 500:
                time.sleep(2 ** attempt)
            else:
                raise
    return ""

# 文字数厳格化
def finalize_ad(text: str, target_chars: int) -> str:
    text = normalize_output(text)

    # 句点保証
    if not text.endswith("。"):
        text += "。"

    if len(text) <= target_chars:
        return text

    # 長い → 文末優先でカット
    cut = text[:target_chars]
    for i in range(len(cut) - 1, -1, -1):
        if cut[i] == "。":
            return cut[:i + 1]

    # 句点がなければ強制
    return cut[:-1] + "。"

# キーワード指定
def build_keyword_instruction(keywords: str):
    words = [w for w in re.split(r"[ 　]+", keywords.strip()) if w]
    if not words:
        return ""
    return f"以下のキーワードを必ずすべて1回以上含めてください。絶対に省略しないでください。文章の自然な位置に挿入してください：" + "、".join(words) + "。"

# 表現
def build_tone_instruction(tone: str) -> str:
    if tone == "やさしい":
        return (
            "・小学生でもすぐ理解できる、やさしい言葉を使う\n"
            "・難しい言葉や専門用語は使わない\n"
            "・一文を短めにして素直な文にする\n"
        )
    else:
        return (
            "・公の場に出しても問題ない表現にする\n"
            "・誤解を招く言い方や強すぎる表現は避ける\n"
        )

# 広告文生成
def generate_newspaper_ad_api(text, target_chars, keywords, tone, temperature=0.2):
    if not HF_TOKEN:
        raise RuntimeError("HUGGINGFACEHUB_API_TOKEN が設定されていません。")

    client = InferenceClient(model=MODEL_ID, token=HF_TOKEN, timeout=60)
    cleaned = remove_strings(text)
    tone_inst = build_tone_instruction(tone)
    keyword_inst = build_keyword_instruction(keywords)

    ad = _call_chat(
        client,
        [{"role": "user", "content": (
            f"次の原稿をもとに、テレビ番組の視聴を促す新聞広告文を作成してください。\n"
            f"・日本語のみ\n"
            f"・固有名詞禁止\n"
            f"・改行なし\n"
            f"・文末は「。」\n"
            f"・文字数は {target_chars + 30} 文字以上でもよい\n\n"
            f"・読者が『この番組を見てみたい』と思うように書く\n"
            f"・番組の魅力を簡潔にまとめる\n"
            f"・最後に視聴を促す一言を入れる\n"
            f"{tone_inst}\n"
            f"{keyword_inst}\n\n"
            f"【原稿】\n{cleaned}\n\n【広告文】"
        )}],
        max_tokens=int((target_chars + 30) * 3),
        temperature=temperature,
    )

    ad = normalize_output(ad)
    ad = finalize_ad(ad, target_chars)
    return ad

# Word保存
def create_docx_bytes(text):
    doc = Document()
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# 履歴保存
def make_history_label(option: str, text: str, uploadfile=None) -> str:
    if option == "ファイル" and uploadfile is not None:
        return f"{uploadfile.name}"
    else:
        head = normalize_output(text)[:10]
        return f"{head}…"

# Streamlit UI
def main():
    load_css("style.css")
    
    if "current_ad" not in st.session_state:
        st.session_state["current_ad"] = ""

    # 履歴保持
    if "history" not in st.session_state:
        st.session_state["history"] = []

    st.title("広告生成")
    
    option = st.sidebar.radio("入力方法を選択", ("テキスト", "ファイル"))
    text = ""

    if option == "テキスト":
        text = st.text_area("広告文にしたい原稿を入力してください", height=260)

    else:
        uploadfile = st.file_uploader("ファイルを選択", type=["txt", "docx"])
        if uploadfile is not None:
            try:
                if uploadfile.name.endswith(".txt"):
                    text = uploadfile.read().decode("utf-8", errors="ignore")
                elif uploadfile.name.endswith(".docx"):
                    doc = Document(uploadfile)
                    text = "\n".join([p.text for p in doc.paragraphs])
                # st.success("ファイルを読み込みました")
            except Exception as e:
                st.error(f"ファイルの読み込みに失敗しました: {e}")
                text = ""

    # 文字数指定
    target_chars = st.sidebar.number_input("文字数", min_value=10, max_value=500, value=120, step=1)
    
    # 文章表現選択
    tone = st.sidebar.radio("文章のスタイル", ["かたい", "やさしい"], horizontal=True)
    
    # キーワード指定
    keywords = st.sidebar.text_input("キーワード指定（スペース区切り）", value="")

    # 保存ファイル名
    filename = st.sidebar.text_input("保存するファイル名", value="newspaper")
    ext = st.sidebar.radio("保存形式", [".txt", ".docx"], horizontal=True)
    download = filename + ext
    
    # 履歴    
    st.sidebar.subheader("履歴")
    for item in st.session_state["history"]:
        with st.sidebar.expander(item["label"]):
            st.write(item["content"])

    # 要約生成    
    if st.button("広告文を生成"):
        try:
            if not text.strip():
                st.warning("原稿を入力してください。")
            else:
                with st.spinner("広告文を生成中..."):
                    ad = generate_newspaper_ad_api(
                        text=text,
                        target_chars=target_chars,
                        keywords=keywords,
                        tone=tone
                    )
                    label = make_history_label(
                        option,
                        text,
                        uploadfile if option == "ファイル" else None
                    )

                    st.session_state["current_ad"] = ad
                    st.session_state["history"].insert(
                        0,
                        {
                            "label": label,
                            "content": ad
                        }
                    )
                    st.session_state["history"] = st.session_state["history"][:5]
       
        except Exception as e:
            st.error(f"エラーが発生しました: {e}")

    # ダウンロードボタン
    if st.session_state["current_ad"]:
        st.text_area("生成された広告文", st.session_state["current_ad"], height=200)
        st.markdown(f"文字数：{len(st.session_state['current_ad'])} 文字")
                
        if ext == ".docx":
            file_data = create_docx_bytes(st.session_state["current_ad"])
            mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            file_data = st.session_state["current_ad"]
            mime = "text/plain"
            
        st.download_button(
            label="ダウンロード",
            data=file_data,
            file_name=download,
            mime=mime
            )

# 実行
if __name__ == "__main__":
    main()
